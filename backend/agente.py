import os, json, requests
import re
import tempfile
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

def get_key(name):
    return os.environ.get(name, "")

def ps(q):
    k = get_key("PARALLEL_API_KEY")
    if not k: return {"results": []}
    try:
        r = requests.post("https://api.parallel.ai/v1/search",
            headers={"Content-Type": "application/json", "x-api-key": k},
            json={"objective": q, "search_queries": [q], "mode": "fast"}, timeout=10)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        print(f"Parallel search error: {e}")
        return {"results": []}

def pe(url, obj=""):
    k = get_key("PARALLEL_API_KEY")
    if not k or not url: return ""
    try:
        r = requests.post("https://api.parallel.ai/v1/extract",
            headers={"Content-Type": "application/json", "x-api-key": k},
            json={"urls": [url], "objective": obj or "Extract"}, timeout=10)
        r.raise_for_status()
        d = r.json()
        ex = d.get("results", [{}])[0].get("excerpts", [""])
        return ex[0][:1500] if ex else ""
    except Exception as e:
        print(f"Parallel extract error: {e}")
        return ""

def gm(prompt, retries=2):
    k = get_key("GEMINI_API_KEY")
    if not k:
        return "Error: GEMINI_API_KEY not configured"
    models_to_try = ["gemini-1.5-flash", "gemini-1.5-pro"]
    for model in models_to_try:
        for attempt in range(retries):
            try:
                base_url = "https://generativelanguage.googleapis.com/v1beta/models"
                r = requests.post(f"{base_url}/{model}:generateContent",
                    json={"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"temperature": 0.5, "maxOutputTokens": 8192}},
                    params={"key": k}, timeout=60)
                if r.status_code == 200:
                    return r.json()["candidates"][0]["content"]["parts"][0]["text"]
                elif r.status_code == 404:
                    break
                elif r.status_code == 429:
                    import time
                    time.sleep(15 * (attempt + 1))
                    continue
            except Exception as e:
                print(f"Gemini error ({model}): {e}")
                break
    return "Error: Gemini API failed"

def reverse_geocode(lat, lng):
    """Convert coordinates to city/neighborhood name using BigDataCloud."""
    try:
        r = requests.get(f"https://api.bigdatacloud.net/data/reverse-geocode-client?latitude={lat}&longitude={lng}&localityLanguage=en", timeout=5)
        if r.status_code == 200:
            data = r.json()
            city = data.get("city") or ""
            state = data.get("principalSubdivision") or ""
            country = data.get("countryName") or ""
            neighborhood = data.get("locality") or ""
            return {"city": city, "neighborhood": neighborhood, "state": state, "country": country,
                    "full": f"{neighborhood}, {city}, {state}, {country}" if neighborhood else f"{city}, {state}, {country}"}
    except:
        pass
    return {"city": "", "neighborhood": "", "state": "", "country": "", "full": f"{lat}, {lng}"}

def get_sag_aftra_rates():
    return {
        "background_actor_daily": 231, "background_actor_hourly": 29,
        "principal_actor_daily": 1246, "principal_actor_weekly": 4357,
        "pension_health_pct": 20.5, "meal_penalty_first": 25,
        "meal_penalty_second": 35, "meal_penalty_subsequent": 50,
        "meal_break_hours": 6, "minimum_background_la": 57,
        "overtime_hourly_multiplier": 1.5, "golden_time_threshold_hours": 12,
        "golden_time_multiplier": 2.0, "rest_period_hours": 12,
        "travel_day_rate": 300, "wardrobe_allowance": 44, "stunt_base_daily": 2608,
    }

def get_insurance_breakdown(crew_size, equipment_value, days):
    rates = {"general_liability_daily": 150, "workers_comp_pct": 0.025,
             "equipment_daily_pct": 0.001, "cast_insurance_daily": 200,
             "errors_omissions_daily": 100, "auto_liability_daily": 75, "umbrella_policy_daily": 250}
    estimated_payroll = crew_size * 400 * days
    breakdown = {
        "general_liability": rates["general_liability_daily"] * days,
        "workers_comp": int(estimated_payroll * rates["workers_comp_pct"]),
        "equipment": int(equipment_value * rates["equipment_daily_pct"] * days) if equipment_value > 0 else 0,
        "cast_insurance": rates["cast_insurance_daily"] * days,
        "errors_omissions": rates["errors_omissions_daily"] * days,
        "auto_liability": rates["auto_liability_daily"] * days,
        "umbrella": rates["umbrella_policy_daily"] * days, "total": 0,
    }
    breakdown["total"] = sum(v for k, v in breakdown.items() if k != "total")
    return breakdown

def get_parking_requirements(crew_size, extras):
    total_people = crew_size + extras
    personal_vehicles = total_people // 2
    production_vehicles = 2 if crew_size <= 10 else 4
    if crew_size > 20: production_vehicles += 2
    return {
        "personal_vehicles": personal_vehicles, "production_vehicles": production_vehicles,
        "total_vehicles": personal_vehicles + production_vehicles,
        "parking_cost_daily": (personal_vehicles + production_vehicles) * 25,
        "truck_cost_daily": production_vehicles * 150,
        "basecamp_needed": crew_size > 15, "basecamp_cost_daily": 500 if crew_size > 15 else 0,
    }

def get_catering_requirements(crew_size, extras, shooting_hours=12):
    total_people = crew_size + extras
    meals_per_day = 2 if shooting_hours >= 10 else 1
    return {
        "total_people": total_people, "meals_per_day": meals_per_day,
        "snack_cost_daily": total_people * 35, "meal_cost_daily": total_people * 45 * meals_per_day,
        "total_daily_catering": (total_people * 45 * meals_per_day) + (total_people * 35),
        "meal_penalty_first_30min": 25, "meal_penalty_second_30min": 35, "meal_penalty_subsequent": 50,
        "turnaround_hours": 6, "golden_time_threshold": 12,
    }

def get_budget_estimate(crew_size, extras, shooting_days=1):
    sag = get_sag_aftra_rates()
    principals = min(5, max(1, extras // 5))
    principals_total = principals * sag["principal_actor_daily"] * shooting_days
    extras_total = extras * sag["background_actor_daily"] * shooting_days
    crew_total = crew_size * 500 * shooting_days
    sag_contribution = int((extras_total + principals_total) * (sag["pension_health_pct"] / 100))
    permit = 1620 if (crew_size + extras) > 30 else 0
    insurance = 1500 * shooting_days
    parking = get_parking_requirements(crew_size, extras)
    parking_total = (parking["parking_cost_daily"] + parking["truck_cost_daily"] + parking["basecamp_cost_daily"]) * shooting_days
    catering = get_catering_requirements(crew_size, extras)
    catering_total = catering["total_daily_catering"] * shooting_days
    equipment_total = (2000 if crew_size <= 10 else 3500) * shooting_days
    total = extras_total + principals_total + crew_total + sag_contribution + permit + insurance + parking_total + catering_total + equipment_total
    return {
        "shooting_days": shooting_days, "principals": principals,
        "principals_daily": principals * sag["principal_actor_daily"], "principals_total": principals_total,
        "extras": extras, "extras_daily": extras * sag["background_actor_daily"], "extras_total": extras_total,
        "crew_daily": crew_size * 500, "crew_total": crew_total,
        "sag_contribution": sag_contribution, "permit": permit, "insurance": insurance,
        "parking_total": parking_total, "catering_total": catering_total, "equipment_total": equipment_total,
        "total": total, "contingency": int(total * 0.1), "grand_total": int(total * 1.1),
    }

def extract_production_info(message):
    msg_lower = message.lower()
    location = None
    location_type = "unknown"
    found_state = None
    countries = []
    
    coord_match = re.search(r'(-?\d+\.?\d*)\s*,\s*(-?\d+\.?\d*)', message)
    if coord_match:
        lat, lng = float(coord_match.group(1)), float(coord_match.group(2))
        location = {"lat": lat, "lng": lng, "name": f"{lat}, {lng}"}
        location_type = "coordinates"
        geo = reverse_geocode(lat, lng)
        if geo["city"]:
            location.update({"city": geo["city"], "state": geo["state"], "neighborhood": geo["neighborhood"], "full_address": geo["full"]})
            # Auto-detect country from geocoding if no country found
            if not countries:
                geo_country = geo.get("country", "")
                if "united states" in geo_country.lower() or "usa" in geo_country.lower():
                    countries.append("United States")
                    if geo.get("state"):
                        found_state = geo["state"]
                elif geo_country:
                    # Map country names
                    country_map = {"mexico": "Mexico", "colombia": "Colombia", "spain": "Spain", "japan": "Japan"}
                    for key, val in country_map.items():
                        if key in geo_country.lower():
                            countries.append(val)
                            break
    
    maps_match = re.search(r'(https?://(?:www\.)?google\.com/maps/[^\s]+)', message)
    if maps_match:
        location = {"url": maps_match.group(1), "name": "Google Maps Location"}
        location_type = "maps_url"
    
    if not location:
        near_match = re.search(r'(?:near|in|at|around)\s+([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)', message)
        if near_match:
            location = {"name": near_match.group(1)}
            location_type = "place_name"
    
    film_countries = {
        "mexico": "Mexico", "colombia": "Colombia", "spain": "Spain",
        "argentina": "Argentina", "brazil": "Brazil", "chile": "Chile",
        "peru": "Peru", "costa rica": "Costa Rica", "japan": "Japan",
        "uk": "United Kingdom", "france": "France", "germany": "Germany",
        "italy": "Italy", "india": "India", "south korea": "South Korea",
        "australia": "Australia", "canada": "Canada", "usa": "United States",
        "czech republic": "Czech Republic", "hungary": "Hungary",
        "thailand": "Thailand", "uae": "United Arab Emirates",
    }
    
    usa_states = {
        "california": "California", "los angeles": "California",
        "hollywood": "California", "san francisco": "California", "san diego": "California",
        "new york": "New York", "brooklyn": "New York", "manhattan": "New York",
        "georgia": "Georgia", "atlanta": "Georgia", "savannah": "Georgia",
        "louisiana": "Louisiana", "new orleans": "Louisiana",
        "new mexico": "New Mexico", "albuquerque": "New Mexico",
        "texas": "Texas", "austin": "Texas", "houston": "Texas", "dallas": "Texas",
        "illinois": "Illinois", "chicago": "Illinois",
        "florida": "Florida", "miami": "Florida", "orlando": "Florida",
        "colorado": "Colorado", "denver": "Colorado",
        "washington": "Washington", "seattle": "Washington",
        "oregon": "Oregon", "portland": "Oregon",
        "tennessee": "Tennessee", "nashville": "Tennessee",
        "arizona": "Arizona", "phoenix": "Arizona",
        "utah": "Utah", "salt lake city": "Utah",
        "nevada": "Nevada", "las vegas": "Nevada",
        "massachusetts": "Massachusetts", "boston": "Massachusetts",
        "pennsylvania": "Pennsylvania", "philadelphia": "Pennsylvania",
        "north carolina": "North Carolina",
        "south carolina": "South Carolina",
        "michigan": "Michigan", "detroit": "Michigan",
        "ohio": "Ohio", "cleveland": "Ohio", "columbus": "Ohio",
        "virginia": "Virginia",
        "maryland": "Maryland", "baltimore": "Maryland",
        "new jersey": "New Jersey",
        "connecticut": "Connecticut",
        "alabama": "Alabama",
        "kentucky": "Kentucky",
        "minnesota": "Minnesota", "minneapolis": "Minnesota",
        "wisconsin": "Wisconsin",
        "indiana": "Indiana",
        "missouri": "Missouri", "kansas city": "Missouri",
        "hawaii": "Hawaii", "honolulu": "Hawaii",
        "alaska": "Alaska",
    }
    
    # Check general countries first
    for key, val in film_countries.items():
        if key in msg_lower and val not in countries:
            countries.append(val)
    
    # Check USA states - require word boundaries for 2-3 letter abbreviations
    state_abbrs = [k for k in usa_states.keys() if len(k) <= 3]
    if state_abbrs:
        state_pattern = r'\b(' + '|'.join(re.escape(k) for k in state_abbrs) + r')\b'
        state_match = re.search(state_pattern, msg_lower)
        if state_match:
            found_state = usa_states[state_match.group(1)]
            if "United States" not in countries:
                countries.append("United States")
    
    # Check full state names (4+ chars) with substring match
    if not found_state:
        for key, val in usa_states.items():
            if len(key) > 3 and key in msg_lower:
                found_state = val
                if "United States" not in countries:
                    countries.append("United States")
                break
    
    if not countries:
        return {"error": True, "message": "Please specify a destination country or US state. Examples: California, New York, Mexico, Spain, Japan, etc."}
    
    scene_type = "urban"
    if any(w in msg_lower for w in ["colonial", "historic", "old town"]):
        scene_type = "heritage"
    elif any(w in msg_lower for w in ["mountain", "beach", "forest", "desert", "nature"]):
        scene_type = "natural"
    elif any(w in msg_lower for w in ["studio", "indoor", "interior", "soundstage"]):
        scene_type = "studio"
    elif any(w in msg_lower for w in ["aerial", "drone", "fly"]):
        scene_type = "aerial"
    elif any(w in msg_lower for w in ["water", "lake", "river", "sea", "ocean", "underwater"]):
        scene_type = "water"
    
    extras = 0
    m = re.search(r'(\d+)\s*extras?\b', msg_lower)
    if m: extras = int(m.group(1))
    
    principals = 0
    m = re.search(r'(\d+)\s*(?:actors?|principals?|talent)', msg_lower)
    if m: principals = int(m.group(1))
    
    crew_size = 10
    m = re.search(r'(\d+)\s*(?:crew|people|person|staff|team)', msg_lower)
    if m: crew_size = int(m.group(1))
    else: crew_size = max(10, extras // 2)
    
    drones = any(w in msg_lower for w in ["drone", "drones", "uav", "quadcopter", "fpv"])
    pyrotechnics = any(w in msg_lower for w in ["pyro", "pyrotechnics", "fireworks", "explosion", "fire", "burn"])
    night_shoot = any(w in msg_lower for w in ["night", "evening", "dusk", "dark"])
    water_related = any(w in msg_lower for w in ["water", "lake", "river", "sea", "ocean", "boat", "ship"])
    
    budget_usd = 0
    budget_patterns = [
        r'(?:budget|cost|spend|invest\s+of)\s*\$?(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*(k|thousand|million|m|usd)?',
        r'\$(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*(k|thousand|million|m)',
        r'(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*(?:million|m)\s*(?:budget|project|film|cost)',
    ]
    for pat in budget_patterns:
        m = re.search(pat, msg_lower)
        if m:
            amount = float(m.group(1).replace(",", ""))
            unit = (m.group(2) or "").lower()
            if unit in ["k", "thousand"]: amount *= 1000
            elif unit in ["million", "m"]: amount *= 1000000
            budget_usd = int(amount)
            break
    
    return {
        "countries": countries, "location": location, "location_type": location_type,
        "scene_type": scene_type, "extras": extras, "principals": principals,
        "crew_size": crew_size, "drones": drones, "pyrotechnics": pyrotechnics,
        "night_shoot": night_shoot, "water_related": water_related,
        "budget_usd": budget_usd, "state": found_state, "error": False
    }

def get_usa_state_info(state):
    states = {
        "California": {
            "film_office": "California Film Commission",
            "incentives": "20-25% tax credit (Film & TV Tax Credit 4.0)",
            "locations": "Hollywood, LA, SF, SD, Santa Barbara",
            "permits": "FilmLA (LA city) + State permits",
            "unions": "SAG-AFTRA, IATSE, Teamsters 399",
            "showstoppers": "FilmLA permits required; High union costs; Parking restrictions",
            "hotels": "The Garland, Magic Castle Hotel, Taglyan Complex",
            "hospitals": "Cedars-Sinai (LA), UCSF (SF), UCSD (SD)",
        },
        "New York": {
            "film_office": "Governor's Office for Motion Picture & TV",
            "incentives": "25-30% tax credit (NY State Film Tax Credit)",
            "locations": "Manhattan, Brooklyn, Queens, Hudson Valley",
            "permits": "MOME (Mayor's Office of Media & Entertainment)",
            "unions": "SAG-AFTRA, IATSE 52, Teamsters 817",
            "showstoppers": "MOME permits required; Congestion pricing; Noise ordinances",
            "hotels": "The Pod Hotel, The Local NYC, The William Vale",
            "hospitals": "NYU Langone, Mount Sinai, Bellevue",
        },
        "Georgia": {
            "film_office": "Georgia Film Office",
            "incentives": "20-30% tax credit (GA Entertainment Industry Act)",
            "locations": "Atlanta, Savannah, Senoia, Crawfordville",
            "permits": "Local city/county + GA Film Office",
            "unions": "SAG-AFTRA, IATSE (lower presence)",
            "showstoppers": "Less infrastructure than CA/NY; Weather delays",
            "hotels": "The Whitley, Embassy Suites Atlanta, The Gastonian",
            "hospitals": "Emory Healthcare (ATL), Memorial Health, Candler (Sav)",
        },
        "Louisiana": {
            "film_office": "Louisiana Entertainment",
            "incentives": "25-40% tax credit (LA Motion Picture Tax Credit)",
            "locations": "New Orleans, Baton Rouge, Shreveport",
            "permits": "LA Film Office + local",
            "unions": "SAG-AFTRA, IATSE (moderate)",
            "showstoppers": "Hurricane season (June-Nov); Humidity affects gear",
            "hotels": "The Roosevelt Windsor Court, Omni Royal Orleans",
            "hospitals": "Ochsner Medical Center, Tulane Medical Center",
        },
        "New Mexico": {
            "film_office": "New Mexico Film Office",
            "incentives": "25-35% tax credit (NM Film Tax Credit)",
            "locations": "Albuquerque, Santa Fe, Las Cruces",
            "permits": "NM Film Office + local",
            "unions": "SAG-AFTRA, IATSE (smaller market)",
            "showstoppers": "High altitude affects gear; Limited crew pool",
            "hotels": "Hotel Parq Central, El Rey Inn, Isleta Resort",
            "hospitals": "UNM Hospital (ABQ), Christus St. Vincent (SF)",
        },
    }
    return states.get(state, {
        "film_office": f"{state} Film Commission/Office",
        "incentives": "Check state film office for current incentives",
        "locations": "Contact state film office",
        "permits": "State + local film permits required",
        "unions": "SAG-AFTRA, IATSE may apply",
        "showstoppers": "Check with state film office",
        "hotels": "Contact local production services",
        "hospitals": "Verify nearest trauma center",
    })

def get_country_info(country):
    info = {
        "Mexico": {"risk": "HIGH", "permit_cost": "$500-$5,000/day", "processing_time": "10-15 days",
                   "restrictions": "No foreign drone ops; AFAC permit mandatory; Extras need permits",
                   "showstoppers": "Drone ban for foreigners; Extra work permits; Heritage restrictions",
                   "insurance": "AXA Mexico, GNP Seguros; Allianz Global for foreign equip; IMSS for locals",
                   "medical": "IMSS/ISSSTE (public); Hospital Angeles, ABC (private)"},
        "Colombia": {"risk": "MEDIUM", "permit_cost": "$300-$3,000/day", "processing_time": "5-10 days",
                     "restrictions": "Film commission approval; Aerocivil drone permits; Extras need visas",
                     "showstoppers": "Visa requirements; Customs delays; Language barriers",
                     "insurance": "SURA, Mapfre; Lloyd's for foreign equip; ARL (mandatory)",
                     "medical": "EPS (public); Hospital Universitario San Ignacio (private)"},
        "Spain": {"risk": "MEDIUM", "permit_cost": "$200-$4,000/day", "processing_time": "10-20 days",
                  "restrictions": "Autonomous region approvals; Heritage restrictions; EU drone regs",
                  "showstoppers": "Autonomy bureaucracy; Spanish bureaucracy; Heritage permits",
                  "insurance": "Mapfre, Allianz Spain; EU equip covered; Mutualidad (mandatory)",
                  "medical": "SNS (public); Hospital Quirón (private); EHIC for EU crew"},
        "Japan": {"risk": "HIGH", "permit_cost": "$1,000-$10,000/day", "processing_time": "14-30 days",
                  "restrictions": "Foreign crew limits; Strict drone regs; Complex location permits",
                  "showstoppers": "Strict foreign crew rules; Complex bureaucracy; High costs",
                  "insurance": "Tokio Marine, Sompo Japan; Local insurance required for foreign equip",
                  "medical": "NHI (public); St. Luke's International (private)"},
        "United States": {"risk": "LOW-MEDIUM", "permit_cost": "$100-$2,000/day", "processing_time": "5-14 days",
                          "restrictions": "State-specific permits; Location releases; Union regs",
                          "showstoppers": "Union pickup fees; Insurance requirements; Location release laws",
                          "insurance": "Film Emissary, AIG Entertainment; Inland marine for equip",
                          "medical": "Mayo Clinic, Cedars-Sinai, NYU Langone"},
    }
    return info.get(country, {"risk": "MEDIUM", "permit_cost": "$200-$4,000/day", "processing_time": "7-21 days",
                               "restrictions": "Local film commission approval needed",
                               "showstoppers": "Unknown local regulations; Permit delays",
                               "insurance": "Local insurance required",
                               "medical": "Verify nearest hospital"})

def get_country_vendors(country):
    vendors = {
        "Mexico": [
            "• [Story Productions](https://story.mx/) — Full service production",
            "• [We Produce](https://weproduce.mx/) — Equipment & crew",
            "• [80 Days Films](https://80daysfilms.com/) — International co-productions",
            "• [Mexico Film Commission](https://www.filmcommission.gob.mx/) — Permits & locations",
        ],
        "United States": [
            "• [Film Emissary](https://www.filmemissary.com/) — Insurance",
            "• [Wrapbook](https://www.wrapbook.com/) — Payroll + Insurance",
            "• [ShareGrid](https://www.sharegrid.com/) — Equipment rental",
            "• [ProductionHUB](https://www.productionhub.com/) — Crew & vendors",
            "• [Mandy](https://www.mandy.com/) — Crew finder",
            "• [Peerspace](https://www.peerspace.com/) — Location rentals",
            "• [Giggster](https://www.giggster.com/) — Film locations",
            "• [SAG-AFTRA](https://www.sagaftra.org/) — Union resources",
            "• [FilmLA](https://www.filmla.com/) — LA permits",
        ],
        "California": [
            "• [FilmLA](https://www.filmla.com/) — LA City permits",
            "• [CA Film Commission](https://www.film.ca.gov/) — State incentives",
            "• [SAG-AFTRA](https://www.sagaftra.org/) — Union rates",
        ],
        "New York": [
            "• [MOME](https://www.nyc.gov/site/mome/index.page) — NYC permits",
            "• [NY Film Commission](https://esd.ny.gov/new-york-state-film-tax-credit-program) — Incentives",
        ],
        "Georgia": [
            "• [GA Film Office](https://www.georgia.org/industries/film-entertainment) — State film office",
        ],
        "Louisiana": [
            "• [LA Entertainment](https://www.louisianaentertainment.gov/) — State film office",
        ],
        "New Mexico": [
            "• [NM Film Office](https://nmfilm.com/) — State film office",
        ],
        "Colombia": [
            "• [ProColombia Film](https://www.procolombia.co/en/industries/creative-industries/film) — Film commission",
            "• [Cartagena Film Festival](https://www.cartagenafilmfestival.com/) — Festival info",
        ],
        "Spain": [
            "• [Spain Film Commission](https://www.spainfilmcommission.com/) — Locations & permits",
        ],
        "Japan": [
            "• [Japan Film Commission](https://www.japanfc.jp/eng/) — Production resources",
        ],
        "United Kingdom": [
            "• [British Film Commission](https://britishfilmcommission.org.uk/) — UK production",
        ],
    }
    return "\n".join(vendors.get(country, ["• Contact local film office for vendor recommendations"]))

def generate_demo_report(data):
    if data.get("error"):
        return f"**{data['message']}**"
    
    cs = ", ".join(data["countries"])
    st = data.get("state", "")
    loc_str = f" ({st})" if st else ""
    
    location_info = ""
    if data.get("location") and data["location"].get("city"):
        city = data["location"]["city"]
        neighborhood = data["location"].get("neighborhood", "")
        state = data["location"].get("state", "")
        location_info = f"\n\n📍 FILMING LOCATION: {neighborhood}{', ' if neighborhood else ''}{city}, {state}"
        location_info += f"\n   Coordinates: {data['location']['lat']}, {data['location']['lng']}"
        location_info += f"\n   Google Maps: https://www.google.com/maps?q={data['location']['lat']},{data['location']['lng']}"
    elif data.get("location") and data["location"].get("name"):
        location_info = f"\n\n📍 FILMING LOCATION: {data['location']['name']}"
    
    report = f"""## Film Production Report
Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

### 1. Executive Summary
Production for {cs}{loc_str} ({data['scene_type']} location) with {data['extras']} extras, {data['principals']} principals, {data['crew_size']} crew.
Budget: ${data['budget_usd']:,}. Key challenges: permits, crew, compliance, insurance.{location_info}

### 2. Country Analysis
"""
    for c in data["countries"]:
        info = get_country_info(c)
        report += f"""
{c.upper()} [{info['risk']}]
├── Permit Cost: {info['permit_cost']}
├── Processing Time: {info['processing_time']}
├── Key Restrictions:
    • {info['restrictions']}
└── Show Stoppers:
    ⚠️ {info['showstoppers']}
"""
    
    if st:
        si = get_usa_state_info(st)
        report += f"""
### 2b. {st} State Details
├── Film Office: {si['film_office']}
├── Incentives: {si['incentives']}
├── Key Locations: {si['locations']}
├── Permits: {si['permits']}
├── Unions: {si['unions']}
├── Show Stoppers: {si['showstoppers']}
└── Hotels: {si['hotels']} | Hospitals: {si['hospitals']}
"""
    
    budget = get_budget_estimate(data['crew_size'], data['extras'])
    sag = get_sag_aftra_rates()
    
    report += f"""
### 3. Detailed Budget Estimate (1 shooting day)
├── Principals ({budget['principals']} actors): ${budget['principals_daily']:,}/day × {budget['principals']} = ${budget['principals_total']:,}
├── Extras ({budget['extras']} people): ${sag['background_actor_daily']}/day × {budget['extras']} = ${budget['extras_total']:,}
├── Crew ({data['crew_size']} people): ${budget['crew_daily']:,}/day
├── SAG Pension & Health ({sag['pension_health_pct']}%): ${budget['sag_contribution']:,}
├── Permit: ${budget['permit']:,}
├── Insurance: ${budget['insurance']:,}
├── Parking/Trucks: ${budget['parking_total']:,}
├── Catering ({data['crew_size']+data['extras']} people): ${budget['catering_total']:,}
├── Equipment Rental: ${budget['equipment_total']:,}
├── SUBTOTAL: ${budget['total']:,}
├── Contingency (10%): ${budget['contingency']:,}
└── GRAND TOTAL (1 day): ${budget['grand_total']:,}
"""

    report += f"""
### 4. SAG-AFTRA Requirements (2025-2026)
├── Background Actors: ${sag['background_actor_daily']}/day (8 hours)
├── Principal Actors: ${sag['principal_actor_daily']}/day
├── Pension & Health: {sag['pension_health_pct']}% of gross
├── Meal Break: Every {sag['meal_break_hours']} hours
├── Meal Penalty (1st 30min): ${sag['meal_penalty_first']}
├── Meal Penalty (2nd 30min): ${sag['meal_penalty_second']}
├── Golden Time (>{sag['golden_time_threshold_hours']}hrs): {sag['golden_time_multiplier']}× hourly
├── Rest Period: {sag['rest_period_hours']} hours between wrap/call
└── Travel Day: ${sag['travel_day_rate']}
"""

    ins = get_insurance_breakdown(data['crew_size'], 50000, 1)
    report += f"""
### 5. Insurance Breakdown (1 day)
├── General Liability ($1M): ${ins['general_liability']}
├── Workers' Compensation: ${ins['workers_comp']}
├── Equipment Coverage: ${ins['equipment']}
├── Cast Insurance: ${ins['cast_insurance']}
├── Errors & Omissions: ${ins['errors_omissions']}
├── Auto Liability: ${ins['auto_liability']}
├── Umbrella Policy: ${ins['umbrella']}
└── TOTAL DAILY INSURANCE: ${ins['total']}
"""

    parking = get_parking_requirements(data['crew_size'], data['extras'])
    catering = get_catering_requirements(data['crew_size'], data['extras'])
    
    report += f"""
### 6. Parking & Logistics
├── Personal Vehicles: {parking['personal_vehicles']} cars
├── Production Vehicles: {parking['production_vehicles']} trucks/vans
├── Total Vehicles: {parking['total_vehicles']}
├── Parking Cost: ${parking['parking_cost_daily']}/day
├── Truck Rental: ${parking['truck_cost_daily']}/day
├── Basecamp Needed: {'Yes' if parking['basecamp_needed'] else 'No'}
├── Basecamp Cost: ${parking['basecamp_cost_daily']}/day
└── TOTAL PARKING/DAY: ${parking['parking_cost_daily'] + parking['truck_cost_daily'] + parking['basecamp_cost_daily']}

### 7. Catering Requirements
├── Total People: {catering['total_people']}
├── Meals/Day: {catering['meals_per_day']}
├── Snack Cost: ${catering['snack_cost_daily']}/day
├── Meal Cost: ${catering['meal_cost_daily']}/day
├── TOTAL CATERING/DAY: ${catering['total_daily_catering']}
├── Meal Penalty (1st 30min): ${catering['meal_penalty_first_30min']}
├── Meal Penalty (2nd 30min): ${catering['meal_penalty_second_30min']}
└── Meal Penalty (after): ${catering['meal_penalty_subsequent']}
"""

    report += """
### 8. Next Steps & Contacts
1. Contact local film office for specific permit requirements
2. Secure insurance quotes from recommended providers
3. Verify property permissions (written authorization)
4. File permit application (3-5 business days before shoot)
5. Reserve parking/basecamp
6. Coordinate SAG-AFTRA casting if using union actors
7. Arrange catering (dietary restrictions, meal timing)
8. Verify nearest hospital with ER + foreign language support
"""

    # Contextual links based on country/state
    country_links = get_country_vendors(data['countries'][0])
    state = data.get("state", "")
    
    if state:
        state_links = get_country_vendors(state)
        report += f"""
### 9. References & Links
{country_links}

{state_links}
"""
    else:
        report += f"""
### 9. References & Links
{country_links}
"""
    
    return report

def process_query(message, demo_mode=False):
    data = extract_production_info(message)
    if data.get("error"):
        return generate_demo_report(data)
    
    use_demo = demo_mode or not get_key("GEMINI_API_KEY")
    if use_demo:
        return generate_demo_report(data)
    
    # Live research - gather real links from Parallel
    search_results = []
    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = {}
        for c in data["countries"]:
            futures[executor.submit(ps, f"{c} film production insurance providers 2025")] = f"{c}_insurance"
            futures[executor.submit(ps, f"{c} film commission permit office website")] = f"{c}_film_office"
            futures[executor.submit(ps, f"{c} film equipment rental companies")] = f"{c}_equipment"
            if data["drones"]:
                futures[executor.submit(ps, f"{c} drone permit regulations filming")] = f"{c}_drone"
            if data["pyrotechnics"]:
                futures[executor.submit(ps, f"{c} pyrotechnics filming permit special effects")] = f"{c}_pyro"
        
        for future in as_completed(futures, timeout=30):
            try:
                result = future.result()
                if result.get("results"):
                    search_results.extend(result["results"][:2])
            except:
                pass
    
    real_links = []
    for r in search_results:
        title = r.get("title", "")
        url = r.get("url", "")
        if url and title:
            real_links.append(f"• [{title}]({url})")
    
    links_text = "\n".join(real_links[:15]) if real_links else "No specific links found in search results."
    
    cs = ", ".join(data["countries"])
    st = data.get("state", "")
    
    prompt = f"""Film production analysis for {cs}{' ('+st+' state)' if st else ''}.
PRODUCTION: {data['scene_type']} location, {data['extras']} extras, {data['principals']} principals, {data['crew_size']} crew, ${data['budget_usd']:,} budget
DRONES: {data['drones']}, PYRO: {data['pyrotechnics']}, NIGHT: {data['night_shoot']}

SEARCH RESULTS FOUND:
{links_text}

WRITE A DETAILED REPORT:
1. Executive Summary
2. Country/State Analysis Table with real data from search
3. Detailed Budget Estimate (with real numbers)
4. SAG-AFTRA Requirements (2025-2026 rates)
5. Insurance Breakdown (medical, equipment, liability)
6. Parking & Logistics (vehicles, costs)
7. Catering Requirements (meal penalties)
8. Next Steps & Contacts
9. References & Links (ONLY use links from search results above - do not invent URLs)

IMPORTANT: Only include URLs that appear in the search results. Do not make up website addresses.
Be specific about showstoppers - what could STOP production."""
    
    result = gm(prompt)
    if result.startswith("Error:"):
        return generate_demo_report(data)
    return result

def create_app():
    from flask import Flask, request, jsonify, send_file
    from flask_cors import CORS
    import io
    
    app = Flask(__name__, static_folder=os.path.join(os.path.dirname(__file__), '..', 'frontend'), static_url_path='')
    CORS(app)
    
    @app.route("/")
    def index(): return send_file(app.static_folder + '/index.html')
    
    @app.route("/api/chat", methods=["POST"])
    def chat():
        data = request.json
        msg = data.get("message", "")
        demo = data.get("demo_mode", False)
        if not msg: return jsonify({"success": False, "error": "Empty"}), 400
        try:
            result = process_query(msg, demo_mode=demo)
            return jsonify({"success": True, "response": result})
        except Exception as e:
            print(f"Chat error: {e}")
            data = extract_production_info(msg)
            return jsonify({"success": True, "response": generate_demo_report(data)})
    
    @app.route("/api/export-docx", methods=["POST"])
    def export_docx():
        data = request.json
        report_text = data.get("report", "")
        if not report_text: return jsonify({"success": False, "error": "No report"}), 400
        try:
            docx_bytes = generate_docx(report_text)
            return send_file(io.BytesIO(docx_bytes),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True, download_name="production-passport-report.docx")
        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500
    
    @app.route("/api/health", methods=["GET"])
    def health():
        return jsonify({"status": "ok", "mode": "live" if get_key("GEMINI_API_KEY") else "demo"})
    return app

def generate_docx(text):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    title = doc.add_heading("Production Passport Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif '☐' in line:
            doc.add_paragraph(line.replace('☐', '').strip()).style = 'List Bullet'
        elif '•' in line:
            clean = line.lstrip('•').strip()
            if clean: doc.add_paragraph(clean).style = 'List Bullet'
        elif any(x in line for x in ['├──', '└──', '│', '[HIGH]', '[MEDIUM]', '[LOW]']):
            doc.add_paragraph(line.replace('├── ', '').replace('└── ', '').replace('│', '').strip())
        elif line and not line.startswith('['):
            doc.add_paragraph(line.replace('**', ''))
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        with open(tmp.name, 'rb') as f:
            data = f.read()
        os.unlink(tmp.name)
    return data

if __name__ == "__main__":
    app = create_app()
    app.run(host="0.0.0.0", port=8080, debug=False)